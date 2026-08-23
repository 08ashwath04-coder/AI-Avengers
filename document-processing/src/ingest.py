"""Extract supported uploads into the active AI knowledge store."""

import csv
import json
import os
import sys

BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
PROJECT_DIR = os.path.dirname(BASE_DIR)
KNOWLEDGE_FILE = os.path.join(PROJECT_DIR, "ai-search", "data", "uploaded_documents.json")
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".csv", ".xls", ".xlsx"}


def text_record(source, text, page=None, sheet=None, row=None):
    record = {"source": source, "text": text.strip()}
    if page is not None:
        record["page"] = page
    if sheet:
        record["sheet"] = sheet
    if row:
        record["row"] = row
    return record


def extract_pdf(file_path, source):
    from pypdf import PdfReader

    records = []
    for page_number, page in enumerate(PdfReader(file_path).pages, start=1):
        content = page.extract_text() or ""
        if content.strip():
            records.append(text_record(source, content, page=page_number))
    return records


def extract_docx(file_path, source):
    from docx import Document

    document = Document(file_path)
    sections = [paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip()]
    for table in document.tables:
        for values in table.rows:
            cells = [cell.text.strip() for cell in values.cells]
            if any(cells):
                sections.append(" | ".join(cells))
    return [text_record(source, "\n\n".join(sections))] if sections else []


def extract_txt(file_path, source):
    with open(file_path, "r", encoding="utf-8-sig", errors="replace") as file:
        content = file.read()
    return [text_record(source, content)] if content.strip() else []


def row_records(source, sheet, rows):
    rows = [list(row) for row in rows if any(value not in (None, "") for value in row)]
    if not rows:
        return []
    headers = [str(value).strip() or f"Column {index + 1}" for index, value in enumerate(rows[0])]
    records = []
    for row_number, values in enumerate(rows[1:], start=2):
        facts = {
            headers[index]: "" if index >= len(values) or values[index] is None else str(values[index]).strip()
            for index in range(len(headers))
        }
        text = f"Sheet: {sheet}\n" + "\n".join(f"{key}: {value}" for key, value in facts.items())
        records.append(text_record(source, text, sheet=sheet, row=facts))
    return records or [text_record(source, f"Sheet: {sheet}\nColumns: {', '.join(headers)}", sheet=sheet)]


def extract_csv(file_path, source):
    with open(file_path, "r", encoding="utf-8-sig", newline="", errors="replace") as file:
        return row_records(source, "CSV", list(csv.reader(file)))


def extract_xlsx(file_path, source):
    from openpyxl import load_workbook

    workbook = load_workbook(file_path, read_only=True, data_only=True)
    records = []
    for worksheet in workbook.worksheets:
        records.extend(row_records(source, worksheet.title, list(worksheet.iter_rows(values_only=True))))
    return records


def extract_xls(file_path, source):
    import xlrd

    workbook = xlrd.open_workbook(file_path)
    records = []
    for worksheet in workbook.sheets():
        records.extend(row_records(source, worksheet.name, [worksheet.row_values(index) for index in range(worksheet.nrows)]))
    return records


def extract(file_path):
    extension = os.path.splitext(file_path)[1].lower()
    source = os.path.basename(file_path)
    if extension not in SUPPORTED_EXTENSIONS:
        raise ValueError("Unsupported file type. Please upload PDF, DOCX, TXT, CSV, XLS, or XLSX.")
    handlers = {
        ".pdf": extract_pdf, ".docx": extract_docx, ".txt": extract_txt,
        ".csv": extract_csv, ".xls": extract_xls, ".xlsx": extract_xlsx,
    }
    return handlers[extension](file_path, source)


def ingest(file_path):
    records = extract(file_path)
    if not records:
        raise ValueError("Unable to extract readable content from this file.")

    os.makedirs(os.path.dirname(KNOWLEDGE_FILE), exist_ok=True)
    existing = []
    if os.path.exists(KNOWLEDGE_FILE):
        with open(KNOWLEDGE_FILE, "r", encoding="utf-8") as file:
            existing = json.load(file)
    source = os.path.basename(file_path)
    existing = [record for record in existing if record.get("source") != source]
    with open(KNOWLEDGE_FILE, "w", encoding="utf-8") as file:
        json.dump(existing + records, file, ensure_ascii=False, indent=2)
    return len(records)


if __name__ == "__main__":
    try:
        if len(sys.argv) != 2:
            raise ValueError("A file path is required.")
        print(json.dumps({"records": ingest(sys.argv[1])}))
    except Exception as error:
        print(json.dumps({"error": str(error)}))
        sys.exit(1)
